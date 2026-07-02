#!/usr/bin/env python3
"""Build bilingual lease agreement .docx (Indonesian left, English right)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = Path(__file__).resolve().parents[1] / 'documents' / 'Perjanjian-Sewa-Menyewa-Lease-Agreement.docx'


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for side, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        margins.append(node)
    tc_pr.append(margins)


def shade_borders(table):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def add_paragraph(cell, text: str, *, bold=False, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if cell.paragraphs and paragraph.text == '' and len(cell.paragraphs) == 1:
        pass
    else:
        paragraph = cell.add_paragraph()
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return paragraph


def set_cell_text(cell, text: str, *, bold=False, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    cell.text = ''
    for index, line in enumerate(text.split('\n')):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = align
        run = paragraph.add_run(line)
        run.bold = bold
        run.underline = underline
        run.font.name = 'Arial'
        run.font.size = Pt(size)


def add_article_header_row(table, id_pasal: str, id_sub: str, en_pasal: str, en_sub: str):
    row = table.add_row()
    for cell, pasal, sub in ((row.cells[0], id_pasal, id_sub), (row.cells[1], en_pasal, en_sub)):
        cell.text = ''
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(pasal)
        r1.bold = True
        r1.font.name = 'Arial'
        r1.font.size = Pt(11)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(sub)
        r2.bold = True
        r2.underline = True
        r2.font.name = 'Arial'
        r2.font.size = Pt(11)
        set_cell_margins(cell)


def add_row(table, indonesian: str, english: str):
    row = table.add_row()
    set_cell_text(row.cells[0], indonesian)
    set_cell_text(row.cells[1], english)
    for cell in row.cells:
        set_cell_margins(cell)


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'nil')
        borders.append(element)
    tbl_pr.append(borders)


def set_row_height(row, height_cm: float):
    tr = row._tr
    tr_pr = tr.get_or_add_trPr()
    tr_height = OxmlElement('w:trHeight')
    tr_height.set(qn('w:val'), str(int(height_cm * 567)))
    tr_height.set(qn('w:hRule'), 'exact')
    tr_pr.append(tr_height)



def add_signature_section(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(18)

    sig_table = doc.add_table(rows=4, cols=2)
    remove_table_borders(sig_table)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for column in sig_table.columns:
        column.width = Cm(8.3)

    header_row = sig_table.rows[0]
    set_cell_text(header_row.cells[0], 'Pihak Pertama / First Party', align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(header_row.cells[1], 'Pihak Kedua / Second Party', align=WD_ALIGN_PARAGRAPH.CENTER)

    signature_row = sig_table.rows[1]
    set_row_height(signature_row, 5.0)
    for cell in signature_row.cells:
        cell.text = ''
        set_cell_margins(cell, top=40, start=120, bottom=40, end=120)

    role_row = sig_table.rows[2]
    set_cell_text(role_row.cells[0], '(Pemilik Properti)', align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(role_row.cells[1], '(Penyewa Properti)', align=WD_ALIGN_PARAGRAPH.CENTER)

    name_row = sig_table.rows[3]
    set_cell_text(name_row.cells[0], '_______________________________', align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(name_row.cells[1], 'Michael William Veitch', align=WD_ALIGN_PARAGRAPH.CENTER)

    date_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(date_table)
    date_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for column in date_table.columns:
        column.width = Cm(8.3)
    set_cell_text(date_table.rows[0].cells[0], 'Date: _______________________', align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(date_table.rows[0].cells[1], 'Date: _______________________', align=WD_ALIGN_PARAGRAPH.CENTER)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('PERJANJIAN SEWA MENYEWA')
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(12)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run('LEASE AGREEMENT')
    run2.bold = True
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    shade_borders(table)
    table.autofit = False
    for column in table.columns:
        column.width = Cm(8.3)

    add_row(
        table,
        'Kami, yang bertanda tangan di bawah ini:',
        'We, the undersigned below:',
    )

    add_row(
        table,
        (
            '1. Nama: __________________________________________________________________\n'
            'No. KTP: __________________________________________________________________\n'
            'Alamat: _________________________________________________________________\n'
            'Selanjutnya disebut sebagai Pihak Pertama (Pemilik)'
        ),
        (
            '1. Name:\n'
            'KTP/ID No.:\n'
            'Address:\n'
            'Hereafter referred to as the First Party (The Lessor)'
        ),
    )

    add_row(
        table,
        (
            '2. Nama: Michael William Veitch\n'
            'No. Paspor: P606835HO\n'
            'Kewarganegaraan: Kanada\n'
            'Selanjutnya disebut sebagai Pihak Kedua (Penyewa)'
        ),
        (
            '2. Name: Michael William Veitch\n'
            'Passport No.: P606835HO\n'
            'Nationality: Canadian\n'
            'Hereafter referred to as the Second Party (The Lessee)'
        ),
    )

    add_row(
        table,
        (
            'Properti yang disewakan adalah sebuah bangunan yang terletak di:\n'
            'Alamat: Jl. Sekar Wangi, Gang III No. 22, Kesiman, Kertalangu, Denpasar Timur, Kota Denpasar, Bali 80237\n'
            'Listrik: 2.300 watt\n'
            'Fasilitas: Satu tempat tidur, Satu unit AC'
        ),
        (
            'The property being rented is a house located at:\n'
            'Address: Jl. Sekar Wangi, Gang III No. 22, Kesiman, Kertalangu, Denpasar Timur, Kota Denpasar, Bali 80237\n'
            'Electric capacity: 2,300 watts\n'
            'Facilities: One bed, One AC unit'
        ),
    )

    add_row(
        table,
        (
            'Pihak Pertama bermaksud untuk menyewakan properti tersebut kepada Pihak Kedua '
            'dan Pihak Kedua setuju untuk menyewa properti tersebut dari Pihak Pertama.'
        ),
        (
            'The First Party intends to rent the property to the Second Party, and the Second Party '
            'agrees to rent the property from the First Party.'
        ),
    )

    add_row(
        table,
        (
            'Pihak Pertama adalah pemilik sah dari properti yang disewakan dan menjamin bahwa Pihak Kedua '
            'tidak akan diganggu atau digugat oleh pihak ketiga mengenai hak mereka sebagai penyewa properti.'
        ),
        (
            'The First Party is the rightful owner of the rented property and guarantees that the Second Party '
            'will not be disturbed or sued by any third party concerning their right as a tenant of the property.'
        ),
    )

    add_row(
        table,
        'Pihak Kedua telah berkunjung dan melihat dan menerima kondisi properti.',
        'The Second Party has visited, seen, and accepted the condition of the property.',
    )

    add_row(
        table,
        'Perjanjian sewa ini telah dibuat dan telah diterima dengan syarat dan ketentuan sebagai berikut:',
        'This lease agreement has been established and has been accepted with an agreement to the terms and conditions as follows:',
    )

    articles = [
        (
            'Pasal 1', 'Masa Sewa', 'Article 1', 'The Period of Lease',
            (
                'Masa sewa terhitung dari tanggal 1 Juli 2026 sampai 30 Juni 2027.\n'
                'Perjanjian ini berakhir secara otomatis pada tanggal 1 Juli 2027 kecuali diperpanjang oleh kedua belah pihak.\n'
                'Jika sewa tidak diperpanjang, Pihak Pertama, pada waktu yang disepakati oleh kedua belah pihak, dapat mengunjungi properti untuk pemeriksaan dan keperluan lainnya.\n'
                'Jika Pihak Pertama tidak ingin memperpanjang sewa, maka Pihak Pertama wajib memberitahukan kepada Pihak Kedua minimal 1 (satu) bulan sebelum tanggal berakhirnya perjanjian ini.'
            ),
            (
                'The period of lease is from July 1, 2026, to June 30, 2027.\n'
                'This agreement will automatically be terminated by July 1, 2027 unless renewed by both parties.\n'
                'If the lease is not renewed, the First Party, at a time agreed upon by both parties, may visit the property for inspection and other purposes.\n'
                'If the First Party does not desire to extend the contract, the First Party shall notify the Second Party at least 1 (one) month prior to the termination date of this agreement.'
            ),
        ),
        (
            'Pasal 2', 'Harga Sewa dan Pembayaran', 'Article 2', 'Rental Rate and Payment',
            (
                'Kedua belah pihak setuju bahwa harga sewa properti selama 1 (satu) tahun adalah IDR 41.000.000 (empat puluh satu juta Rupiah).\n'
                'Pembayaran sewa telah diselesaikan secara penuh pada tanggal 24 Juni 2026 melalui transfer bank.'
            ),
            (
                'Both parties have agreed that the rental price for the property for 1 (one) year is IDR 41,000,000 (forty one million Rupiah).\n'
                'The rental payment has been completed in full on 24 June, 2026 via bank transfer.'
            ),
        ),
        (
            'Pasal 3', 'Perpanjangan Sewa', 'Article 3', 'Lease Extension',
            (
                'Jika Pihak Kedua ingin memperpanjang masa sewa, pemberitahuan kepada Pihak Pertama dilakukan 1 (satu) bulan sebelum tanggal berakhirnya perjanjian ini.\n'
                'Pembayaran untuk memperpanjang sewa harus dibayar maksimal 3 (tiga) hari setelah sewa selesai.'
            ),
            (
                'If the Second Party wishes to extend the rental period, notice must be given to the First Party 1 (one) month before the expiration date of this agreement.\n'
                'The payment for the extended lease period must be paid a maximum of 3 (three) days after the current period ends.'
            ),
        ),
        (
            'Pasal 4', 'Tujuan Sewa', 'Article 4', 'The Purpose of Lease',
            'Properti harus digunakan oleh Pihak Kedua sebagai tempat tinggal.',
            'The property must be used by The Second Party strictly for residential purposes.',
        ),
        (
            'Pasal 5', 'Larangan', 'Article 5', 'Prohibitions',
            (
                'Pihak Kedua dilarang menyewakan properti kepada pihak ketiga.\n'
                'Pihak Kedua dilarang melakukan penambahan atau perubahan pada bangunan tanpa persetujuan tertulis dari Pihak Pertama.\n'
                'Pihak Kedua dilarang menggunakan properti untuk aktivitas yang melanggar hukum, termasuk namun tidak terbatas pada:\n'
                'a. Aktivitas kriminal\n'
                'b. Perdagangan atau penggunaan narkoba\n'
                'c. Kegiatan perjudian ilegal\n'
                'd. Pelacuran atau eksploitasi seksual\n'
                'e. Penyimpanan atau distribusi barang-barang terlarang\n'
                'f. Kegiatan lain yang melanggar hukum yang berlaku'
            ),
            (
                'The Second Party is prohibited from renting out the property to a third party (subletting).\n'
                'The Second Party is prohibited from making additions or structural changes to the building without the First Party’s written approval.\n'
                'The Second Party is prohibited from using the property for illegal activities, including but not limited to:\n'
                'a. Criminal activities\n'
                'b. Drug trade or use\n'
                'c. Illegal gambling activities\n'
                'd. Prostitution or sexual exploitation\n'
                'e. Storage or distribution of prohibited goods\n'
                'f. Other activities that violate applicable laws'
            ),
        ),
        (
            'Pasal 6', 'Tanggung Jawab Penyewa', 'Article 6', 'The Lessee Responsibilities',
            (
                'Pihak Kedua bertanggung jawab atas pemeliharaan dan kebersihan bangunan.\n'
                'Kerusakan pada bangunan atau furnitur harus diberitahukan kepada Pihak Pertama dan diperbaiki.\n'
                'Pihak Kedua bertanggung jawab melakukan pemeliharaan terhadap mesin air, AC, seluruh furnitur dan barang elektronik lainnya untuk menghindari terjadinya kerusakan pada fasilitas properti sehingga pada akhir masa sewa dapat dikembalikan dalam kondisi lengkap dan baik.\n'
                'Pihak Kedua bertanggung jawab atas pembelian listrik prabayar.'
            ),
            (
                'The Second Party is responsible for the maintenance and cleanliness of the building.\n'
                'Any damage to the building or furniture must be reported to the First Party and repaired.\n'
                'The Second Party is responsible for the upkeep of the water pump, air conditioning units, furniture, and other electronic devices to avoid any damage to the property facilities, ensuring that at the end of the lease period, they can be returned in complete and good condition.\n'
                'The Second Party is responsible for the purchase of prepaid electricity tokens.'
            ),
        ),
        (
            'Pasal 7', 'Tanggung Jawab Pemilik', 'Article 7', 'The Lessor Responsibilities',
            (
                'Pihak Pertama bertanggung jawab atas pemeliharaan struktural bangunan seperti: atap bocor, dinding retak, kegagalan instalasi pipa, dan kegagalan instalasi listrik yang menyebabkan tidak layak huni terhadap properti.'
            ),
            (
                'The First Party is responsible for the structural maintenance of the building, such as: leaky roofs, cracked walls, pipe installation failures, and electrical wiring failures that render the property uninhabitable.'
            ),
        ),
        (
            'Pasal 8', 'Keadaan Terpaksa (Force Majeure)', 'Article 8', 'Force Majeure',
            (
                'Apabila terjadi kerugian atau kerusakan terhadap properti disebabkan oleh keadaan memaksa di luar kemampuan para pihak, misalnya: bencana alam seperti banjir, gempa bumi, angin topan, kerusuhan massa atau bencana alam lainnya, para pihak sepakat untuk mendiskusikan kerugian yang ditanggung para pihak, dan mengambil solusi yang adil bagi para pihak atas kerugian.'
            ),
            (
                'In the event of loss or damage to the property caused by circumstances of force majeure beyond the ability of the parties—for example, natural disasters such as floods, earthquakes, hurricanes, mass riots, or other unpredictable disasters—the parties agree to discuss losses mutually and take a fair solution for both parties for the losses.'
            ),
        ),
        (
            'Pasal 9', 'Resolusi Perselisihan', 'Article 9', 'Dispute Resolution',
            (
                'Jika Pihak Pertama membatalkan perjanjian ini, maka Pihak Pertama harus mengembalikan sisa uang sewa secara proporsional.\n'
                'Jika Pihak Kedua membatalkan perjanjian ini, maka uang sewa tidak akan dikembalikan.\n'
                'Setiap perselisihan yang timbul dari pelaksanaan perjanjian ini akan diselesaikan dengan musyawarah dan mufakat bersama. Namun jika penyelesaian bersama tersebut gagal mencapai konsensus apapun, maka kedua pihak sepakat untuk memilih domisili hukum tetap dan umum di Kantor Pengadilan Negeri Denpasar.'
            ),
            (
                'If the First Party cancels this agreement early, the First Party must refund the remaining rental money proportionally.\n'
                'If the Second Party cancels this agreement early, the paid rental money will not be refunded.\n'
                'Any disputes arising from the implementation of this agreement will be resolved through mutual discussion and consensus. However, if such a resolution fails to reach any consensus, both parties agree to choose the fixed and general legal domicile at the Denpasar District Court.'
            ),
        ),
        (
            'Pasal 10', 'Penutup', 'Article 10', 'Closing',
            (
                'Perpanjangan sewa setelah habis masa sewa, para pihak setuju membuat perjanjian baru dengan membuat kesepakatan harga oleh para pihak.\n'
                'Pada akhir sewa tanpa perpanjangan, Pihak Kedua mengembalikan properti beserta furnitur di dalamnya seperti semula kepada Pihak Pertama tanpa dihuni/diduduki oleh siapapun, dalam kondisi baik.\n'
                'Perjanjian sewa ini dibuat tanpa paksaan dari pihak lain dan dibuat dalam 2 (dua) salinan di atas materai dan ditandatangani oleh kedua belah pihak dan memiliki kekuatan hukum.\n'
                'Segala hal yang tidak atau belum diatur dalam perjanjian ini akan diselesaikan melalui diskusi bersama antara para pihak.\n'
                'Jika dalam perjanjian ini ditemukan pemahaman yang bertentangan antara bahasa Indonesia dan bahasa Inggris, maka pemahaman yang berlaku adalah pemahaman dalam bahasa Indonesia.'
            ),
            (
                'Upon the expiration of the lease, the parties agree to create a new agreement for the extension of the lease, including the agreed-upon price.\n'
                'At the end of the lease without an extension, the Second Party must return the property along with its inventory to the First Party in the same condition as it was initially, unoccupied and in good condition.\n'
                'This lease agreement is made without coercion from any party and is created in two (2) copies, stamped and signed by both parties, and carries equal legal force.\n'
                'Any matters not covered or not yet regulated in this agreement will be resolved through mutual discussion between the parties.\n'
                'If any conflicting or ambiguous understanding is found between the Indonesian and English versions of this agreement, the understanding that shall prevail is the one in Indonesian.'
            ),
        ),
    ]

    for id_pasal, id_sub, en_pasal, en_sub, id_body, en_body in articles:
        add_article_header_row(table, id_pasal, id_sub, en_pasal, en_sub)
        add_row(table, id_body, en_body)

    add_signature_section(doc)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f'Wrote {OUTPUT}')


if __name__ == '__main__':
    main()
